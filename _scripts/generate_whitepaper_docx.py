#!/usr/bin/env python3
"""
Generate complete white paper DOCX from modular sections using Pandoc + python-docx.
Works with restructured project (sections/*.md → final DOCX).
"""

import os
import re
import subprocess
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === CONFIGURATION ===
BASE_DIR = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals"
SECTIONS_DIR = os.path.join(BASE_DIR, "sections")
OUTPUT_DIR = os.path.join(BASE_DIR, "_output")
FIGURES_DIR = os.path.join(BASE_DIR, "figures")  # Symlink to _figures/exports
TIMESTAMP = datetime.now().strftime("%Y%m%d-%H-%M")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, f"Nutraceuticals_Whitepaper_{TIMESTAMP}.docx")
TEMP_MD = os.path.join(OUTPUT_DIR, "temp_combined.md")

REPORT_MASTER = os.path.join(BASE_DIR, "report_master.md")
REGISTRY_XLSX = os.path.join(BASE_DIR, "_registry", "source_registry.xlsx")

DEFAULT_SOURCE_NAME_MAP = {
    "S015": "EU/US Nutraceutical Regulatory Framework",
    "S085": "US vs EU Regulatory Dataset",
    "S086": "Regulatory Timeline Dataset",
    "S089": "Whitepaper Master Data Workbook",
    "S104": "Grand View Research 2024 (Modeled Dataset)",
    "S105": "Euromonitor 2024 (Modeled Dataset)",
    "S106": "NBJ 2023 (Modeled Dataset)",
    "S107": "Future Market Insights 2024 (Modeled Dataset)",
    "S108": "MarketsandMarkets 2023 (Modeled Dataset)",
    "S109": "FEDIAF Facts & Figures",
    "S110": "APPA Survey (Modeled Dataset)",
    "S111": "FAO SOFIA",
    "S112": "Eurostat Livestock Dataset",
    "S113": "Mordor Intelligence (Modeled Dataset)",
    "S114": "Nicotra et al. (2025)",
    "S115": "Zoetis 10-K / Filing",
    "S116": "Internal PE/VC Portfolio Mapping",
    "S117": "DSM-Firmenich Annual Report",
    "S118": "Swedencare Annual / Transaction Context",
    "S119": "Virbac Annual Report",
    "S120": "Dechra Annual Report",
    "S121": "Global Antigravity Landscape Composite",
    "S122": "Phytogenic ROI Article Pack",
    "S123": "Urban/Suburban Pet Habits Article",
    "S124": "MARA 194 Regulatory Pack",
    "S125": "Sector Deal Multiples Pack",
    "S126": "EU Green Claims Directive Pack",
    "S127": "Nutrigenomics Review Pack",
    "S128": "Legacy v18 Reference Archive",
}


def load_source_name_map():
    """Load source_id -> short_name from registry with safe fallback."""
    source_map = dict(DEFAULT_SOURCE_NAME_MAP)
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        return source_map

    if not os.path.exists(REGISTRY_XLSX):
        return source_map

    try:
        wb = load_workbook(REGISTRY_XLSX, data_only=True)
        if "Sources" not in wb.sheetnames:
            return source_map
        ws = wb["Sources"]
        headers = [cell.value for cell in ws[1]]
        if "source_id" not in headers or "short_name" not in headers:
            return source_map
        sid_idx = headers.index("source_id")
        name_idx = headers.index("short_name")
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row:
                continue
            sid = str(row[sid_idx] or "").strip()
            short_name = str(row[name_idx] or "").strip()
            if sid and short_name:
                source_map[sid] = short_name
    except Exception:
        return source_map

    return source_map


def format_source_reference_block(raw_block, source_map):
    """
    Convert source block content into human-readable source names plus source IDs.
    Example:
      "S089, Tab: Figure 18; S115" ->
      "Whitepaper Master Data Workbook [S089, Tab: Figure 18]; Zoetis 10-K / Filing [S115]"
    """
    parts = [part.strip() for part in raw_block.split(";") if part.strip()]
    normalized_parts = []
    for part in parts:
        # Expand compact lists like "S115, S116" into two references.
        if re.fullmatch(r"S\d{3}(?:\s*,\s*S\d{3})+", part):
            normalized_parts.extend([token.strip() for token in part.split(",") if token.strip()])
        else:
            normalized_parts.append(part)

    formatted_parts = []
    for part in normalized_parts:
        if "UNVERIFIED" in part.upper():
            formatted_parts.append("Unverified [UNVERIFIED]")
            continue
        match = re.search(r"(S\d{3})", part)
        if not match:
            formatted_parts.append(f"[{part}]")
            continue
        sid = match.group(1)
        short_name = source_map.get(sid, sid)
        formatted_parts.append(f"{short_name} [{part}]")
    return "; ".join(formatted_parts)


def expand_source_lines(markdown_content, source_map):
    """Expand markdown source lines with human-readable source names."""
    output_lines = []
    for line in markdown_content.splitlines():
        stripped = line.strip()
        if stripped.startswith("*Source:") and "[" in stripped and "]" in stripped:
            start = stripped.find("[")
            end = stripped.rfind("]")
            if start != -1 and end != -1 and end > start:
                inner = stripped[start + 1:end]
                expanded = format_source_reference_block(inner, source_map)
                output_lines.append(f"*Source: {expanded}*")
                continue
        output_lines.append(line)
    return "\n".join(output_lines)

def load_section_files_from_master():
    """Load section order from report_master.md."""
    section_files = []
    if not os.path.exists(REPORT_MASTER):
        raise FileNotFoundError(f"Missing report master file: {REPORT_MASTER}")

    with open(REPORT_MASTER, 'r', encoding='utf-8') as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line or line.startswith('#'):
                continue
            if line.startswith('sections/') and line.endswith('.md'):
                section_files.append(os.path.basename(line))

    if not section_files:
        raise ValueError("No section files found in report_master.md")
    return section_files

# Color scheme for styling
COLORS = {
    'primary': RGBColor(0, 48, 87),      # Navy
    'secondary': RGBColor(0, 137, 207),   # Blue
    'accent': RGBColor(208, 74, 2),       # Orange
    'text': RGBColor(51, 51, 51),         # Dark grey
    'grey': RGBColor(139, 155, 165),      # Grey
}

def combine_sections():
    """Combine all section files into single markdown."""
    section_files = load_section_files_from_master()
    source_map = load_source_name_map()
    print(f"Combining {len(section_files)} sections...")
    
    combined_content = []
    
    for section_file in section_files:
        section_path = os.path.join(SECTIONS_DIR, section_file)
        if not os.path.exists(section_path):
            print(f"  ⚠ Skipping missing section: {section_file}")
            continue
        
        with open(section_path, 'r', encoding='utf-8') as f:
            content = f.read()
        content = expand_source_lines(content, source_map)
        
        # Add page break between main sections (not front matter)
        if section_file != "00_front_matter.md":
            combined_content.append("\n\\newpage\n\n")
        
        combined_content.append(content)
        combined_content.append("\n\n")
        print(f"  ✓ {section_file}")
    
    # Write combined markdown
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    with open(TEMP_MD, 'w', encoding='utf-8') as f:
        f.write(''.join(combined_content))
    
    print(f"  ✓ Combined markdown: {TEMP_MD}")
    return TEMP_MD

def convert_with_pandoc(input_md, output_docx):
    """Convert markdown to DOCX using Pandoc."""
    print("Running Pandoc conversion...")
    
    cmd = [
        'pandoc',
        input_md,
        '-o', output_docx,
        '--from', 'markdown+smart+pipe_tables+strikeout',
        '--to', 'docx',
        '--toc',
        '--toc-depth=3',
        '--standalone',
        '--resource-path', f"{SECTIONS_DIR}:{BASE_DIR}",
        '-V', 'mainfont=Georgia',
        '-V', 'sansfont=Arial',
        '--reference-doc=' + os.path.join(BASE_DIR, '_scripts/reference.docx') if os.path.exists(os.path.join(BASE_DIR, '_scripts/reference.docx')) else ''
    ]
    
    # Remove empty reference-doc arg if file doesn't exist
    cmd = [arg for arg in cmd if arg]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"  ✗ Pandoc error: {result.stderr}")
        return False
    
    print(f"  ✓ Pandoc conversion complete")
    return True

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def apply_styling(docx_path):
    """Apply professional styling to the document."""
    print("Applying professional styling...")
    
    doc = Document(docx_path)
    
    h1_count = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        
        # Heading 1 (Part titles)
        if style_name == 'Heading 1':
            h1_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = COLORS['accent']
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            if h1_count > 1:
                para.paragraph_format.page_break_before = True
        
        # Heading 2 (Major sections)
        elif style_name == 'Heading 2':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
        
        # Heading 3
        elif style_name == 'Heading 3':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = COLORS['secondary']
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        
        # Body text
        elif style_name in ['Normal', 'Body Text', 'First Paragraph']:
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS['text']
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(8)
        
        # Figure captions
        if para.text.strip().startswith('Figure') and len(para.text) < 200:
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = COLORS['grey']
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)
    
    # Style tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, '003057')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows (alternating)
        for row_idx, row in enumerate(table.rows[1:], start=1):
            for cell in row.cells:
                if row_idx % 2 == 1:
                    set_cell_shading(cell, 'f4f6f8')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
    
    print(f"  ✓ Styled {h1_count} main sections, {len(doc.tables)} tables")
    doc.save(docx_path)
    return True

def add_cover_page(docx_path):
    """Add professional cover page."""
    print("Adding cover page...")
    
    doc = Document(docx_path)
    
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        
        # Date/Classification
        new_para = first_para.insert_paragraph_before()
        run = new_para.add_run("FEBRUARY 2026 | STRATEGIC INTELLIGENCE")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['grey']
        run.font.bold = True
        new_para.paragraph_format.space_after = Pt(36)
        
        # Title
        title_para = first_para.insert_paragraph_before()
        run = title_para.add_run("ANIMAL NUTRACEUTICALS")
        run.font.name = 'Arial'
        run.font.size = Pt(42)
        run.font.bold = True
        run.font.color.rgb = COLORS['primary']
        
        # Subtitle
        sub_para = first_para.insert_paragraph_before()
        run = sub_para.add_run("Market Bifurcation and Value Creation")
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.color.rgb = COLORS['secondary']
        sub_para.paragraph_format.space_after = Pt(24)
        
        # Tagline
        tag_para = first_para.insert_paragraph_before()
        run = tag_para.add_run("A $13B global market split between\nPet humanization and livestock efficiency")
        run.font.name = 'Georgia'
        run.font.size = Pt(14)
        run.font.italic = True
        tag_para.paragraph_format.space_after = Pt(60)
    
    doc.save(docx_path)
    print("  ✓ Cover page added")

def main():
    print("=" * 60)
    print("NUTRACEUTICALS WHITE PAPER - DOCX GENERATOR")
    print("=" * 60)

    # Step 0: Rebuild figure PNG assets from figures_data.xlsx (source of truth).
    figure_build_script = os.path.join(BASE_DIR, "_scripts", "build_figures_from_excel.py")
    if os.path.exists(figure_build_script):
        print("Regenerating figure assets from master Excel...")
        fig_res = subprocess.run([sys.executable, figure_build_script], capture_output=True, text=True)
        if fig_res.returncode != 0:
            print("ERROR: Figure rebuild failed before DOCX generation.")
            if fig_res.stdout:
                print(fig_res.stdout)
            if fig_res.stderr:
                print(fig_res.stderr)
            return
        print("  ✓ Figure asset regeneration complete")
    else:
        print(f"ERROR: Missing figure build script: {figure_build_script}")
        return
    
    # Step 1: Combine sections
    combined_md = combine_sections()
    
    # Step 2: Convert with Pandoc
    if not convert_with_pandoc(combined_md, OUTPUT_DOCX):
        print("ERROR: Pandoc conversion failed")
        return
    
    # Step 3: Apply styling
    apply_styling(OUTPUT_DOCX)
    
    # Step 4: Add cover page
    add_cover_page(OUTPUT_DOCX)
    
    print("=" * 60)
    print(f"✓ COMPLETE: {OUTPUT_DOCX}")
    print("=" * 60)

if __name__ == "__main__":
    main()
