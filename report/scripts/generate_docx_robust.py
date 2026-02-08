#!/usr/bin/env python3
"""
Robust Markdown to DOCX Converter using Pandoc + python-docx post-processing.
Uses pandoc for proper markdown parsing and python-docx for style enhancements.
"""

import os
import subprocess
import re
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === CONFIGURATION ===
BASE_DIR = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals"
INPUT_MD = os.path.join(BASE_DIR, "report/master_report/Master_WhitePaper_Final.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "report/master_report/Master_WhitePaper_Final.docx")
TEMP_DOCX = os.path.join(BASE_DIR, "report/master_report/temp_pandoc_output.docx")
FIGURES_DIR = os.path.join(BASE_DIR, "report/master_report/figures")
REVIEW_DIR = "/Users/celinecollin/Documents/review"

# Color Scheme
COLORS = {
    'primary': RGBColor(0, 48, 87),      # Navy #003057
    'secondary': RGBColor(0, 137, 207),   # Blue #0089cf
    'accent': RGBColor(208, 74, 2),       # Orange #d04a02
    'text': RGBColor(51, 51, 51),         # Dark grey #333
    'grey': RGBColor(139, 155, 165),      # Grey
}

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    """Set Equity Research style borders: Thick top/bottom, thin inside."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Check if tblBorders exists, if not create
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    # Top Border (Thick Navy)
    for tag in ['w:top', 'w:bottom', 'w:left', 'w:right', 'w:insideH', 'w:insideV']:
        existing = tblBorders.find(qn(tag))
        if existing is not None:
             tblBorders.remove(existing)

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12') # 1.5pt
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '003057')
    tblBorders.append(top)
    
    # Bottom Border (Thick Navy)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '003057')
    tblBorders.append(bottom)
    
    # Inside Horizontal (Thin Grey)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4') # 0.5pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'CCCCCC')
    tblBorders.append(insideH)
    
    # Left/Right/InsideV (None for standard Clean look)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'nil')
    tblBorders.append(left)
    
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'nil')
    tblBorders.append(right)
    
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'nil') # No vertical borders for cleaner look
    tblBorders.append(insideV)

def preprocess_markdown(md_content):
    """Preprocess markdown for better pandoc conversion."""
    lines = md_content.split('\n')
    processed = []
    
    for line in lines:
        stripped = line.strip()
        
        # Skip [INSERT FIGURE] placeholders
        if '[INSERT FIGURE' in stripped:
            continue
        
        # Skip raw URLs on their own lines
        if stripped.startswith('http://') or stripped.startswith('https://'):
            if len(stripped) > 50 and ' ' not in stripped:
                continue
        
        # === COMPREHENSIVE CLEANUP ===
        
        # Clean up source references like source.name+2​ or source.name
        line = re.sub(r'[a-zA-Z-]+\.[a-zA-Z-]+\+\d+\u200b?', '', line)
        
        # Clean up citation brackets [1], [2][3], etc.
        line = re.sub(r'\[\d+\]', '', line)
        
        # Clean orphaned brackets at end of sentences
        line = re.sub(r'\s*\]\s*$', '', line)
        line = re.sub(r'\s*\[\s*$', '', line)
        
        # Clean orphaned brackets in middle of text
        line = re.sub(r'\s+\]', '', line)
        line = re.sub(r'\[\s+', '', line)
        
        # Clean trailing source references like "text.source+2"
        line = re.sub(r'\.[a-zA-Z-]+\+\d+', '.', line)
        
        # Remove zero-width spaces and other invisible characters
        line = line.replace('\u200b', '')
        line = line.replace('\u200c', '')
        line = line.replace('\u200d', '')
        line = line.replace('\ufeff', '')
        
        # Clean up multiple spaces
        line = re.sub(r'  +', ' ', line)
        
        # Clean up space before punctuation
        line = re.sub(r'\s+([.,;:!?])', r'\1', line)
        
        stripped = line.strip()
        
        # === HEADING CONVERSION ===
        
        # Part level: "I. Title" → "# I. Title"
        part_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.\s+(.+)$', stripped)
        if part_match and len(stripped) < 80 and not stripped.startswith('#'):
            processed.append(f'# {stripped}')
            continue
        
        # Section level: "I.1 Title" → "## I.1 Title"
        section_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.?\s+(.+)$', stripped)
        if section_match and len(stripped) < 100 and not stripped.startswith('#'):
            processed.append(f'## {stripped}')
            continue
        
        # Subsection: "I.1.2 Title" → "### I.1.2 Title"
        subsec_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.(\d+)\.?\s+(.+)$', stripped)
        if subsec_match and len(stripped) < 120 and not stripped.startswith('#'):
            processed.append(f'### {stripped}')
            continue
        
        # Sub-subsection: "I.1.2.3 Title" → "#### I.1.2.3 Title"
        subsubsec_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$', stripped)
        if subsubsec_match and len(stripped) < 150 and not stripped.startswith('#'):
            processed.append(f'#### {stripped}')
            continue
        
        processed.append(line)
    
    return '\n'.join(processed)

def convert_with_pandoc(input_md, output_docx):
    """Use pandoc for robust markdown to docx conversion."""
    print("Running pandoc conversion...")
    
    # Read and preprocess markdown
    with open(input_md, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Read cover page
    cover_path = os.path.join(os.path.dirname(input_md), "cover_page.md")
    if os.path.exists(cover_path):
        with open(cover_path, 'r', encoding='utf-8') as f:
            cover_content = f.read()
        content = cover_content + "\n" + content
    
    processed_content = preprocess_markdown(content)
    
    # Write preprocessed content to temp file
    temp_md = input_md.replace('.md', '_preprocessed.md')
    # Add explicit page breaks before H1 headers in preprocessed content
    final_md_content = re.sub(r'\n# ', r'\n\n<div style="page-break-after: always;"></div>\n\n# ', processed_content)
    # Actually, pandoc won't respect that div for docx easily. 
    # We will do it in python-docx post-processing instead.
    with open(temp_md, 'w', encoding='utf-8') as f:
        f.write(processed_content)
    
    # Run pandoc with proper resource paths for images
    cmd = [
        'pandoc',
        temp_md,
        '-o', output_docx,
        '--from', 'markdown+smart+pipe_tables+strikeout+yaml_metadata_block', # No implicit_figures to avoid duplicates
        '--to', 'docx',
        '--toc',
        '--toc-depth=3',
        '--standalone',
        '--resource-path', os.path.dirname(input_md),
        '-V', 'mainfont=Georgia',
        '-V', 'sansfont=Arial',
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Pandoc error: {result.stderr}")
        return False
    
    print(f"  ✓ Pandoc conversion complete")
    
    # Clean up temp file
    # If using debug mode, maybe don't delete?
    # if os.path.exists(temp_md):
    #     os.remove(temp_md)
    return True

def apply_style_enhancements(docx_path):
    """Apply style enhancements to pandoc output."""
    print("Applying style enhancements...")
    
    doc = Document(docx_path)
    
    # Track heading counts for styling
    h1_count = 0
    h2_count = 0
    h3_count = 0
    h4_count = 0
    h5_count = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        
        # Heading 1 styling
        if style_name == 'Heading 1':
            h1_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = COLORS['accent']
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            # Add page break before H1 (except the very first ones on cover)
            if h1_count > 1:
                para.paragraph_format.page_break_before = True
        
        # Heading 2 styling
        elif style_name == 'Heading 2':
            h2_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
        
        # Heading 3 styling
        elif style_name == 'Heading 3':
            h3_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = COLORS['secondary']
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        
        # Heading 4 styling
        elif style_name == 'Heading 4':
            h4_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = COLORS['grey']
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(4)
        
        # Heading 5 styling
        elif style_name == 'Heading 5':
            h5_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = COLORS['grey']
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)
        
        # Body text styling
        elif style_name in ['Normal', 'Body Text', 'First Paragraph', 'Compact', 'Block Quote']:
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS['text']
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(8)
        
        # List styling
        elif 'List' in style_name:
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS['text']

        # === FIGURE ANCHORING ===
        # If paragraph has an image (drawing/blip), keep with next (caption)
        if '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml:
             para.paragraph_format.keep_with_next = True
             para.paragraph_format.space_after = Pt(6)
        
        # If paragraph is a Caption, keep with previous (the image) - actually standard Word handles this if image is above.
        # But we ensure Captions are distinct
        if style_name.startswith('Caption') or (para.text.startswith('Figure') and len(para.text) < 200):
            for run in para.runs:
                 run.font.name = 'Arial'
                 run.font.size = Pt(9)
                 run.font.italic = True
                 run.font.color.rgb = COLORS['grey']
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)
    
    # Style tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style header row (first row)
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, '003057')  # Navy
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Style data rows (alternating)
        for row_idx, row in enumerate(table.rows[1:], start=1):
            for cell in row.cells:
                if row_idx % 2 == 1:
                    set_cell_shading(cell, 'f4f6f8')  # Light grey
                for para in cell.paragraphs:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
        # Autofit table - this is tricky in python-docx, but we can set widths
        table.autofit = True
        set_table_borders(table)
    
    print(f"  ✓ Styled {h1_count} H1, {h2_count} H2, {h3_count} H3, {h4_count} H4, {h5_count} H5 headings")
    print(f"  ✓ Styled {len(doc.tables)} tables")
    
    # === FINAL ARTIFACT CLEANUP ===
    # Remove stray brackets that sometimes appear in headers (e.g. [Heading])
    print("Removing stray brackets...")
    cleaned_count = 0
    for para in doc.paragraphs:
        stripped = para.text.strip()
        # If paragraph starts with [ and ends with ] (or close to it)
        if stripped.startswith('[') and (stripped.endswith(']') or ']' in stripped):
            # Check if it's a heading-like paragraph or Key Finding
            if para.style.name.startswith('Heading') or len(stripped) < 100:
                # Remove the brackets while preserving run formatting where possible
                # Simple text replacement for now as bracket is likely in first/last run
                new_text = stripped.lstrip('[').rstrip(']').strip()
                para.text = new_text
                cleaned_count += 1
        
        # Specific fix for [INSERT FIGURE] remnants or stray [[ ]]
        if '[[' in para.text or ']]' in para.text:
            para.text = para.text.replace('[[', '').replace(']]', '')
            cleaned_count += 1
    
    print(f"  ✓ Removed brackets from {cleaned_count} paragraphs")

    doc.save(docx_path)
    return True

def add_cover_and_toc(docx_path):
    """Add professional cover page and enhance TOC."""
    print("Adding cover page...")
    
    doc = Document(docx_path)
    
    # Get first section
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Insert cover content at the very beginning
    # We'll add a new paragraph at the start
    
    # Find first paragraph
    first_para = doc.paragraphs[0] if doc.paragraphs else None
    
    if first_para:
        # Insert classification before first paragraph
        new_para = first_para.insert_paragraph_before()
        run = new_para.add_run("JANUARY 2026 | STRATEGIC INTELLIGENCE")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['grey']
        run.font.bold = True
        new_para.paragraph_format.space_after = Pt(36)
        
        # Title
        title_para = first_para.insert_paragraph_before()
        run = title_para.add_run("ANIMAL NUTRACEUTICALS")
        run.font.name = 'Arial'
        run.font.size = Pt(42)
        run.font.bold = True
        run.font.color.rgb = COLORS['primary']
        
        # Subtitle
        sub_para = first_para.insert_paragraph_before()
        run = sub_para.add_run("The Wellness Market at an Inflection Point")
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.color.rgb = COLORS['secondary']
        sub_para.paragraph_format.space_after = Pt(24)
        
        # Tagline
        tag_para = first_para.insert_paragraph_before()
        run = tag_para.add_run("Mapping value creation across a $6 billion global industry\ndriven by pet humanization and the post-antibiotic transition")
        run.font.name = 'Georgia'
        run.font.size = Pt(14)
        run.font.italic = True
        tag_para.paragraph_format.space_after = Pt(60)
        
        # Page break after cover
        break_para = first_para.insert_paragraph_before()
        break_para.add_run().add_break()
    
    doc.save(docx_path)
    doc.save(docx_path)
    print("  ✓ Cover page added")

def remove_bookmarks_from_xml(docx_path):
    """Directly manipulate OOXML to remove bookmark tags which cause bracket artifacts."""
    print("Scrubbing bookmark tags from XML...")
    import zipfile
    import shutil
    
    # Create temp directory
    temp_dir = docx_path + "_temp_unzip"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    # Unzip
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)
    
    # Process document.xml
    doc_xml_path = os.path.join(temp_dir, 'word/document.xml')
    if os.path.exists(doc_xml_path):
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
        
        # Remove bookmark start and end tags
        # <w:bookmarkStart w:id="0" w:name="section"/>
        # <w:bookmarkEnd w:id="0"/>
        
        # We use regex for this specific cleanup as it's safer than XML parsing for pure removal without restructuring
        xml_content = re.sub(r'<w:bookmarkStart[^>]*/>', '', xml_content)
        xml_content = re.sub(r'<w:bookmarkStart[^>]*>', '', xml_content) # In case it's not self-closing (rare)
        xml_content = re.sub(r'<w:bookmarkEnd[^>]*/>', '', xml_content)
        
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(xml_content)
    
    # Repack
    # We need to zip the contents of temp_dir back into docx_path
    shutil.make_archive(docx_path.replace('.docx', ''), 'zip', temp_dir)
    shutil.move(docx_path.replace('.docx', '.zip'), docx_path)
    
    # Initial cleanup
    shutil.rmtree(temp_dir)
    print("  ✓ Bookmarks removed from XML")

def main():
    print("=" * 60)
    print("ROBUST MARKDOWN TO DOCX CONVERTER")
    print("=" * 60)
    
    # Step 1: Convert with pandoc
    if not convert_with_pandoc(INPUT_MD, OUTPUT_DOCX):
        print("ERROR: Pandoc conversion failed")
        return
    
    # Step 2: Apply style enhancements
    # Skipped to avoid bracket artifacts (relying on Pandoc default styles)
    apply_style_enhancements(OUTPUT_DOCX)
    
    # Step 3: Add cover page
    add_cover_and_toc(OUTPUT_DOCX)
    
    # Step 4: Scrub bookmarks (artifacts)
    remove_bookmarks_from_xml(OUTPUT_DOCX)
    
    # Step 5: Copy to review folder
    print("Copying to review folder...")
    try:
        os.makedirs(REVIEW_DIR, exist_ok=True)
        review_path = os.path.join(REVIEW_DIR, os.path.basename(OUTPUT_DOCX))
        shutil.copy2(OUTPUT_DOCX, review_path)
        print(f"  ✓ Copied to {review_path}")
    except Exception as e:
        print(f"  ✗ Failed to copy to review folder: {e}")

    print("=" * 60)
    print(f"✓ COMPLETE: {OUTPUT_DOCX}")
    print("=" * 60)

if __name__ == "__main__":
    main()
