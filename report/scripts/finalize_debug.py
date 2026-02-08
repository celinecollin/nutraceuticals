
import os
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_cover_and_toc(docx_path):
    print(f"Adding cover page to {docx_path}...")
    doc = Document(docx_path)
    
    # Create a new document for the cover to ensure it's first
    cover_doc = Document()
    
    # 1. Company Logo/Name
    p = cover_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("NUTRA-INSIGHTS RESEARCH")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 48, 87) # Navy
    
    # Spacer
    for _ in range(5):
        cover_doc.add_paragraph()
        
    # 2. Main Title
    p = cover_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Animal Nutraceuticals:\nThe 2030 Strategic Roadmap")
    run.font.name = 'Arial'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 48, 87) # Navy
    
    # 3. Subtitle
    p = cover_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("From Commodity Feed to Precision Health")
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102) # Grey
    
    # Spacer
    for _ in range(10):
        cover_doc.add_paragraph()
        
    # 4. Date & Version
    p = cover_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("February 2026 | Private Client Edition")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    # Page Break
    cover_doc.add_page_break()
    
    # 5. Table of Contents Placeholder
    # Note: Generating a real TOC in python-docx is hard (requires OXML). 
    # We will add a placeholder or simple text list if we parsed headers.
    # For now, let's just let Word generate it or skip it if user didn't explicitly ask for a *generated* TOC structure in script.
    # The prompt asked for "Declarative Headings" but mostly the DOC should be compliant.
    # We will assume existing DOC has content.
    
    # Merge existing content
    # This is tricky without losing styles. 
    # Better approach: Insert cover elements at the beginning of 'doc'
    
    # Re-open doc to insert at index 0?
    # python-docx doesn't easily support "insert at start".
    # We will use the 'composer' approach or just iterate?
    # Actually, let's just try to insert paragraphs at index 0 of the body.
    
    body = doc._body
    
    # Reverse order insertion to keep them at top
    # Page Break
    p = doc.add_paragraph()
    p.add_run().add_break(type=7) # Page Break
    # We need to move this paragraph to the beginning?
    # body._element.insert(0, p._p)
    
    # Let's try a simpler way: Just save the cover_doc content? 
    # No, merging docs is hard.
    
    # Allow me to try inserting into 'doc' directly at the start.
    
    # Insert Date/Version
    p = doc.add_paragraph("February 2026 | Private Client Edition")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(10)
    body._element.insert(0, p._p)
    
    # Spacer
    for _ in range(10):
         p = doc.add_paragraph()
         body._element.insert(0, p._p)
         
    # Subtitle
    p = doc.add_paragraph("From Commodity Feed to Precision Health")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.runs[0]
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    body._element.insert(0, p._p)
    
    # Main Title
    p = doc.add_paragraph("Animal Nutraceuticals:\nThe 2030 Strategic Roadmap")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 48, 87)
    body._element.insert(0, p._p)

    # Spacer
    for _ in range(5):
         p = doc.add_paragraph()
         body._element.insert(0, p._p)

    # Logo
    p = doc.add_paragraph("NUTRA-INSIGHTS RESEARCH")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 48, 87)
    body._element.insert(0, p._p)
    
    # Page Break after cover?
    # The user content starts now. We should insert a page break at the end of cover.
    # The last inserted element (which is visually first) was 'Logo'.
    # The content follows. We need a page break between "Date" and "Content".
    # The "Date" paragraph is now index ~18.
    # Let's insert a page break paragraph after Date.
    # Wait, 'insert(0)' puts it at top. So 'Logo' is top.
    # We did: Date (top), Spacers (top), Subtitle (top), Title (top), Spacers (top), Logo (top).
    # So order in doc is: Logo, Spacers, Title, Subtitle, Spacers, Date, [Original Content].
    # We want a page break after Date.
    
    # Find the 'Date' paragraph (it's the one we inserted first, which is now deep?)
    # No, 'insert(0)' pushes down.
    # Date was inserted first -> became P[0].
    # Spacers inserted -> Date became P[10].
    # ...
    # Logo inserted -> Logo is P[0].
    # So Date is roughly P[N].
    # We just need to insert a Page Break *after* the Date paragraph.
    # Actually, simpler: Insert a Page Break at 0 *before* we start inserting the cover.
    # Then the Original content is pushed down by the Page Break.
    
    # Let's restart logic.
    pass

def add_cover_v2(docx_path):
    doc = Document(docx_path)
    body = doc._body
    
    # 1. Insert Page Break (separates Cover from Content)
    p_break = OxmlElement('w:p')
    r_break = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_break.append(br)
    p_break.append(r_break)
    body._element.insert(0, p_break)
    
    # 2. Insert Date (visually bottom of cover)
    p = doc.add_paragraph("February 2026 | Private Client Edition")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(10)
    body._element.insert(0, p._p)

    # 3. Spacers
    for _ in range(10):
         p = doc.add_paragraph()
         body._element.insert(0, p._p)

    # 4. Subtitle
    p = doc.add_paragraph("From Commodity Feed to Precision Health")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.runs[0]
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    body._element.insert(0, p._p)
    
    # 5. Title
    p = doc.add_paragraph("Animal Nutraceuticals:\nThe 2030 Strategic Roadmap")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 48, 87)
    body._element.insert(0, p._p)
    
    # 6. Spacers
    for _ in range(5):
         p = doc.add_paragraph()
         body._element.insert(0, p._p)
         
    # 7. Logo (Top)
    p = doc.add_paragraph("NUTRA-INSIGHTS RESEARCH")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 48, 87)
    body._element.insert(0, p._p)
    
    doc.save(docx_path)
    print("Added cover page.")

if __name__ == "__main__":
    add_cover_v2("debug_direct.docx")
