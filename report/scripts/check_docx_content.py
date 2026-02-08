from docx import Document
import sys

def check_docx(path):
    print(f"Checking {path}...")
    try:
        doc = Document(path)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        for i, para in enumerate(doc.paragraphs[:30]):
            text = para.text.strip()
            if text:
                print(f"[{i}] {text[:100]}")
                # Check for brackets at start
                if '[' in text:
                     print(f"   >>> FOUND BRACKET: {text}")
                     print(f"   HEX: {text.encode('utf-8').hex()}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        path = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals/report/master_report/20260118_Master_WhitePaper_Final.docx"
    check_docx(path)
