
import os
import re
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import uuid

# Configuration
INPUT_FILE = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals/report/master_report/Master_WhitePaper_Final.docx"
OUTPUT_FILE = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals/report/master_report/Master_WhitePaper_Final_Visualized.docx"
TEMP_IMG_DIR = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals/report/master_report/figures/temp_viz"

if not os.path.exists(TEMP_IMG_DIR):
    os.makedirs(TEMP_IMG_DIR)

class KeepSafeProtocol:
    def __init__(self, path):
        print(f"Loading document: {path}")
        self.doc = Document(path)

    def save(self, path):
        print(f"Saving document: {path}")
        self.doc.save(path)

class DataMiner:
    def Extract(self, text):
        # Improved Pattern 1: Flexible Money
        # Handles: $13 billion, $13+ billion, $18-24 billion, $1.5-2.0B
        # Groups: 1=StartVal, 2=EndVal, 3=Unit
        money_pattern = r"\$(\d+(?:\.\d+)?)[+]?(?:[–-](\d+(?:\.\d+)?))?\s*(billion|million|B|M)"
        # Note: Text often uses en-dashes or hyphens.
        
        matches = list(re.finditer(money_pattern, text, re.IGNORECASE))
        
        data_points = []
        if len(matches) >= 2:
            for i, m in enumerate(matches):
                val = float(m.group(1))
                if m.group(2): # Handle range
                    val = (val + float(m.group(2))) / 2
                
                # Normalize unit
                unit = m.group(3).lower()
                if 'm' in unit and 'b' not in unit:
                    val = val / 1000.0 # Convert to Billions for standard scale if mixed? 
                    # Or keep as is. Let's assume consistent within a sentence or normalize to value.
                    # For simplicity in this Viz, just value.
                    pass
                
                label = f"Item {i+1}"
                data_points.append({"Label": label, "Value": val})
            
            return pd.DataFrame(data_points), "Growth"

        # Pattern 2: Multiples "15x ... 8x"
        multiple_pattern = r"(\d+(?:\.\d+)?)[+]?(?:[–-](\d+(?:\.\d+)?)?)?x\s*EBITDA"
        matches_mult = list(re.finditer(multiple_pattern, text, re.IGNORECASE))
        
        if len(matches_mult) >= 2:
            data_points = []
            for i, m in enumerate(matches_mult):
                val = float(m.group(1))
                if m.group(2):
                    val = (val + float(m.group(2))) / 2
                label = f"Item {i+1}"
                data_points.append({"Label": label, "Value": val})
            return pd.DataFrame(data_points), "Multiples"

        return None, None

class VizEngine:
    def Generate(self, df, type_name):
        filename = f"{TEMP_IMG_DIR}/{uuid.uuid4()}.png"
        plt.figure(figsize=(6, 4))
        
        colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
        
        plt.bar(df['Label'], df['Value'], color=colors[:len(df)])
        
        if type_name == "Growth":
            plt.title("Financial Projection")
            plt.ylabel("Value")
        elif type_name == "Multiples":
            plt.title("Valuation Multiples (x EBITDA)")
            plt.ylabel("Multiple")
            
        plt.tight_layout()
        plt.savefig(filename, dpi=100)
        plt.close()
        return filename

class GlobalReIndexer:
    def Process(self, doc):
        print("Re-indexing figures...")
        counter = 1
        for para in doc.paragraphs:
            # Flexible strip to handle * or whitespace
            clean_text = para.text.strip().lstrip("*")
            if clean_text.startswith("Figure"):
                # Regex search on the clean text
                # We want to keep the caption part (Group 3)
                match = re.match(r"(Figure\s+)([\w\.]+)(.*)", clean_text)
                if match:
                    caption = match.group(3)
                    # reconstruct
                    new_text = f"Figure {counter}{caption}"
                    if para.text.strip().startswith("*"):
                       new_text = f"*{new_text}" # Restore distinct formatting if it was vaguely markdown-ish
                    
                    para.text = new_text
                    counter += 1
        print(f"Re-indexed {counter-1} figures.")

def main():
    protocol = KeepSafeProtocol(INPUT_FILE)
    miner = DataMiner()
    viz = VizEngine()
    
    insertions = 0
    
    # Iterate through paragraphs
    # Note: modifying list while iterating can be tricky if we insert paragraphs, 
    # but inserting runs/pictures into existing paragraphs is safer, 
    # OR we insert a new paragraph after the current one. 
    # python-docx doesn't easily support "insert paragraph at index".
    # Implementation trick: 
    # We can't easily insert *between* paragraphs in a simple loop without getting lost or breaking structure.
    # However, `paragraph.insert_paragraph_before()` exists. 
    # But we want to insert *after*.
    # Strategy: detailed double-pass or careful index management?
    # Actually, simpler: Collect all insertions first (paragraph_obj, image_path), then apply them.
    
    actions = []
    
    for i, para in enumerate(protocol.doc.paragraphs):
        text = para.text
        if len(text) < 20: continue # Skip short lines
        
        df, type_name = miner.Extract(text)
        if df is not None and len(df) > 0:
            print(f"Found data at para {i}: {type_name}")
            img_path = viz.Generate(df, type_name)
            actions.append((para, img_path))
            
    # Apply insertions
    # To insert *after* a paragraph P, we can find the *next* paragraph and insert *before* it.
    # If P is the last one, we append.
    for para, img_path in actions:
        # Create a new paragraph for the image
        # Determining position: 
        # python-docx insertion is relative. 
        # We can try appending a run to the *current* paragraph with a break? 
        # User requested: "Insert... immediately after the paragraph".
        # The cleanest DOM way is to insert a new paragraph after the current one.
        # But python-docx only has `insert_paragraph_before`.
        # So we find the *next* sibling? python-docx Paragraphs don't know their siblings easily.
        # Workaround: `para._p.addnext(new_p._p)` logic is complex XML manipulation.
        # Alternative: Add run to *end* of current paragraph with a newline? 
        # No, that puts image inline with text.
        # Safe standard way: Iterate and reconstruct? No, destructive.
        
        # XML Insert Solution (Robust)
        # We can use the paragraph's parent (body) and insert new element after.
        new_para = protocol.doc.add_paragraph()
        # Move this new para to right position
        para._p.addnext(new_para._p)
        
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        insertions += 1
        
    print(f"Inserted {insertions} charts.")
    
    # Re-index
    indexer = GlobalReIndexer()
    indexer.Process(protocol.doc)
    
    protocol.save(OUTPUT_FILE)
    print("Done.")

if __name__ == "__main__":
    main()
