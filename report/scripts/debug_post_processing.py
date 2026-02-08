
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Define colors
COLORS = {
    'primary': RGBColor(0, 48, 87),    # Navy Blue
    'secondary': RGBColor(0, 112, 192), # Medium Blue
    'accent': RGBColor(233, 78, 27),    # Orange/Red
    'text': RGBColor(34, 34, 34),       # Dark Grey
    'grey': RGBColor(102, 102, 102)     # Light Grey
}

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)

def set_table_borders(table):
    """Set Equity Research style borders: Thick top/bottom, thin inside."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Check if tblBorders exists, if not create
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    # Top Border (Thick Navy)
    for tag in ['w:top', 'w:bottom', 'w:left', 'w:right', 'w:insideH', 'w:insideV']:
        existing = tblBorders.find(qn(tag))
        if existing is not None:
             tblBorders.remove(existing)

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12') # 1.5pt
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '003057')
    tblBorders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '003057')
    tblBorders.append(bottom)
    
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4') # 0.5pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'CCCCCC')
    tblBorders.append(insideH)
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'nil')
    tblBorders.append(left)
    
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'nil')
    tblBorders.append(right)
    
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'nil')
    tblBorders.append(insideV)

def apply_style_enhancements(docx_path):
    print(f"Applying styles to {docx_path}...")
    doc = Document(docx_path)
    
    h1_count = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        
        if style_name == 'Heading 1':
            h1_count += 1
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = COLORS['accent']
        
        elif style_name == 'Heading 2':
             for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']

        elif style_name in ['Normal', 'Body Text']:
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS['text']
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Figure Anchoring
        if '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml:
             para.paragraph_format.keep_with_next = True
        
        if style_name.startswith('Caption') or (para.text.startswith('Figure') and len(para.text) < 200):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, '003057')
                for para in cell.paragraphs:
                    for run in para.runs:
                         run.font.color.rgb = RGBColor(255,255,255)
                         run.font.bold = True
        
        # Rows
        for row_idx, row in enumerate(table.rows[1:], start=1):
             for cell in row.cells:
                 if row_idx % 2 == 1:
                     set_cell_shading(cell, 'f4f6f8')
        
        table.autofit = True
        set_table_borders(table)

    doc.save(docx_path)
    print("Optimization complete.")

if __name__ == "__main__":
    apply_style_enhancements("debug_direct.docx")
