
import re
import os
import sys
from docx import Document

# Mappings for safe US -> UK conversion
# Combining user instructions and snippet
us_to_uk = {
    "analyze": "analyse", "analysis": "analysis", # Analysis is same? Maybe 'analyzes' -> 'analyses'
    "analyzes": "analyses", "analyzed": "analysed", "analyzing": "analysing",
    "behavior": "behaviour", "behaviors": "behaviours",
    "color": "colour", "colors": "colours",
    "center": "centre", "centers": "centres", "centered": "centred",
    "defense": "defence", 
    "labor": "labour", 
    "organize": "organise", "organizes": "organises", "organized": "organised", "organizing": "organising",
    "organization": "organisation", "organizations": "organisations",
    "realize": "realise", "realizes": "realises", "realized": "realised", "realizing": "realising",
    "program": "programme", "programs": "programmes",
    "optimization": "optimisation", "optimize": "optimise", "optimized": "optimised",
    "standardization": "standardisation", "standardize": "standardise", "standardized": "standardised",
    "capitalization": "capitalisation", "capitalize": "capitalise",
    "urbanization": "urbanisation",
    "specialization": "specialisation", "specialize": "specialise",
    "visualization": "visualisation", "visualize": "visualise"
}

def replace_text_in_run(text):
    if not text:
        return text
        
    # 1. Financial Abbreviation (Regex for precision)
    # Matches "100 million" or "100 Million" -> "100m"
    # Ensure it follows a digit
    text = re.sub(r'(\d)\s+[Mm]illion(s)?', r'\1m', text)
    text = re.sub(r'(\d)\s+[Bb]illion(s)?', r'\1bn', text)
    
    # 2. British English Dictionary Swap
    for us, uk in us_to_uk.items():
        # strict word boundary check to avoid partial replacements
        # Handle lower case
        pattern = r'\b' + re.escape(us) + r'\b'
        text = re.sub(pattern, uk, text)
        
        # Handle Title Case
        pattern_cap = r'\b' + re.escape(us.capitalize()) + r'\b'
        text = re.sub(pattern_cap, uk.capitalize(), text)
        
    return text

def process_element(element):
    """Iterate over runs in a paragraph to preserve formatting."""
    for run in element.runs:
        original = run.text
        replaced = replace_text_in_run(original)
        if original != replaced:
            run.text = replaced

def main():
    input_dir = "report/master_report"
    input_filename = "Master_WhitePaper_Final.docx"
    input_path = os.path.join(input_dir, input_filename)
    
    # Allow running from root
    if not os.path.exists(input_path):
        input_path = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals/report/master_report/Master_WhitePaper_Final.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: Input file not found at {input_path}")
        sys.exit(1)
        
    print(f"Processing {input_path}...")
    doc = Document(input_path)
    
    # Iterate through paragraphs (Body)
    for para in doc.paragraphs:
        process_element(para)

    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_element(para)
                    
    # Note: python-docx doesn't easily expose headers/footers in generic iteration, 
    # but user requested Body, Table Cells, Figure Captions (usually in body paragraphs).
    
    output_filename = "Whitepaper_v4_UK_Formatted.docx"
    output_path = os.path.join(input_dir, output_filename)
    
    doc.save(output_path)
    print(f"Success! Saved formatted document to: {output_path}")

if __name__ == "__main__":
    main()
