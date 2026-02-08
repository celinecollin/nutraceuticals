"""
Publication-Quality DOCX Generator for Animal Nutraceuticals White Paper
Uses python-docx for full control over styling and formatting.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === CONFIGURATION ===
BASE_DIR = "/Users/celinecollin/Library/CloudStorage/OneDrive-Personal/Nutraceuticals"
INPUT_MD = os.path.join(BASE_DIR, "report/master_report/20260118_Master_WhitePaper_Clean.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "report/master_report/20260118_Master_WhitePaper_Polished.docx")
FIGURES_DIR = os.path.join(BASE_DIR, "report/master_report/figures")

# Color Scheme (Matching Data Centres WP)
COLORS = {
    'primary': RGBColor(0, 48, 87),      # Navy Blue #003057
    'secondary': RGBColor(0, 137, 207),  # Bright Blue #0089cf
    'accent': RGBColor(208, 74, 2),      # Burnt Orange #d04a02
    'text': RGBColor(51, 51, 51),        # Dark Grey #333333
    'grey': RGBColor(100, 100, 100),     # Mid Grey
}

# === HELPER FUNCTIONS ===

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    doc.add_page_break()

def create_styles(doc):
    """Define custom styles for the document."""
    styles = doc.styles
    
    # --- Title Style ---
    style = styles.add_style('WP Title', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(36)
    style.font.bold = True
    style.font.color.rgb = RGBColor(255, 255, 255)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # --- Subtitle Style ---
    style = styles.add_style('WP Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(18)
    style.font.italic = True
    style.font.color.rgb = RGBColor(200, 200, 200)
    
    # --- Heading 1 (Section) ---
    style = styles.add_style('WP Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(24)
    style.font.bold = True
    style.font.color.rgb = COLORS['accent']
    style.paragraph_format.space_before = Pt(24)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.keep_with_next = True
    
    # --- Heading 2 (Subsection) ---
    style = styles.add_style('WP Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(18)
    style.font.bold = True
    style.font.color.rgb = COLORS['primary']
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(6)
    
    # --- Heading 3 ---
    style = styles.add_style('WP Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = COLORS['secondary']
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    
    # --- Body Text ---
    style = styles.add_style('WP Body', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS['text']
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- Bullet List ---
    style = styles.add_style('WP Bullet', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS['text']
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.space_after = Pt(4)
    
    # --- Figure Caption ---
    style = styles.add_style('WP Caption', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.italic = True
    style.font.color.rgb = COLORS['grey']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(18)
    
    # --- Blockquote ---
    style = styles.add_style('WP Quote', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Georgia'
    style.font.size = Pt(12)
    style.font.italic = True
    style.font.color.rgb = COLORS['primary']
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    
    # --- Section Divider Number ---
    style = styles.add_style('WP Section Number', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(72)
    style.font.bold = True
    style.font.color.rgb = COLORS['accent']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(150)
    style.paragraph_format.space_after = Pt(24)
    
    # --- Section Divider Title ---
    style = styles.add_style('WP Section Title', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(24)
    style.font.bold = True
    style.font.color.rgb = COLORS['primary']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(200)

def add_section_divider(doc, section_num, section_title):
    """Add a full-page section divider."""
    add_page_break(doc)
    
    # Large section number
    p = doc.add_paragraph(section_num, style='WP Section Number')
    
    # Section title
    p = doc.add_paragraph(section_title.upper(), style='WP Section Title')
    
    add_page_break(doc)

def add_cover_page(doc):
    """Create a full-bleed cover page."""
    # Add a section for the cover
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Date/Classification
    p = doc.add_paragraph()
    run = p.add_run("JANUARY 2026 | STRATEGIC INTELLIGENCE")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS['grey']
    run.font.bold = True
    p.paragraph_format.space_after = Pt(48)
    
    # Main Title
    p = doc.add_paragraph()
    run = p.add_run("ANIMAL NUTRACEUTICALS")
    run.font.name = 'Arial'
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = COLORS['primary']
    p.paragraph_format.space_after = Pt(0)
    
    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run("The Wellness Market at an Inflection Point")
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.font.color.rgb = COLORS['secondary']
    p.paragraph_format.space_after = Pt(24)
    
    # Tagline
    p = doc.add_paragraph()
    run = p.add_run("Mapping value creation across a $6 billion global industry\ndriven by pet humanization and the post-antibiotic transition")
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = COLORS['text']
    p.paragraph_format.space_after = Pt(72)
    
    # Orange Accent Bar (simulated with table)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Inches(2)
    set_cell_shading(cell, 'd04a02')  # Orange
    cell.paragraphs[0].add_run(" ")  # Empty content
    
    # Spacer
    for _ in range(8):
        doc.add_paragraph()
    
    # Footer Classification
    p = doc.add_paragraph()
    run = p.add_run("CONFIDENTIAL | FOR INTERNAL REVIEW")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS['grey']
    run.font.all_caps = True
    
    add_page_break(doc)

def add_executive_summary(doc):
    """Add the executive summary section."""
    p = doc.add_paragraph("EXECUTIVE SUMMARY", style='WP Heading 1')
    
    summary_text = """The global animal nutraceutical market stands at a critical inflection point, valued at approximately USD 6 billion in 2024 and projected to exceed USD 10 billion by 2030. This growth is propelled by two powerful, converging megatrends: the inexorable humanization of companion animals in developed economies, and the regulatory-driven transition away from antibiotic growth promoters in livestock production.

This white paper provides a comprehensive analysis of the market structure, competitive dynamics, and investment thesis across both companion animal and livestock segments. We identify the key value pools, map the competitive landscape from ingredient suppliers to consumer brands, and highlight the strategic implications for investors and operators.

Our analysis reveals a "barbell" market structure: at one end, premium pet supplement brands command EBITDA margins of 20–25% through brand equity and veterinary endorsement; at the other, commodity ingredient suppliers operate on thin margins of 5–10%. The highest risk-adjusted returns lie in IP-protected ingredient platforms and vertically integrated players capturing multiple value chain layers.

The European and North American markets remain the value leaders, while Asia-Pacific—particularly China—represents the primary growth engine. The regulatory environment, particularly the EU's stringent feed additive framework and the global phase-out of antibiotic growth promoters, creates structural tailwinds for probiotic and functional feed solutions.

For investors, the sector offers compelling characteristics: recession-resilient demand (pet spending is notably inelastic), regulatory moats, and consolidation opportunities as the market professionalizes. However, entrants must navigate channel fragmentation, evidentiary requirements for efficacy claims, and the competitive pressure from integrated pharmaceutical players expanding into preventive care."""

    p = doc.add_paragraph(summary_text, style='WP Body')
    
    add_page_break(doc)

def add_table_of_contents(doc):
    """Add a styled table of contents."""
    p = doc.add_paragraph("TABLE OF CONTENTS", style='WP Heading 1')
    
    toc_items = [
        ("I.", "Definition, Scope and Structural Dynamics", 4),
        ("II.", "Functional Segmentation and Clinical Validation", 12),
        ("III.", "Market Structure and Value Capture", 24),
        ("IV.", "Competitive Landscape Analysis", 52),
        ("V.", "Investment Thesis and Transactions", 68),
    ]
    
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num} {title}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = COLORS['primary']
        run.font.bold = True
        
        # Add tab and page number
        tab_run = p.add_run(f"\t{page}")
        tab_run.font.name = 'Arial'
        tab_run.font.size = Pt(12)
        tab_run.font.color.rgb = COLORS['grey']
        
        p.paragraph_format.space_after = Pt(8)
    
    add_page_break(doc)

def is_table_row(line):
    """Check if line is a markdown table row."""
    line = line.strip()
    return line.startswith('|') and line.endswith('|') and line.count('|') >= 2

def is_table_separator(line):
    """Check if line is a markdown table separator (|---|---|)."""
    line = line.strip()
    if not line.startswith('|'):
        return False
    # Check if it contains mostly dashes and pipes
    cleaned = line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
    return len(cleaned) == 0

def parse_table_row(line):
    """Parse a table row into cells."""
    cells = line.strip().split('|')
    # Remove empty first and last elements from split
    cells = [c.strip() for c in cells if c.strip()]
    return cells

def add_markdown_table(doc, table_lines):
    """Convert markdown table lines to a Word table."""
    if len(table_lines) < 2:
        return
    
    # Parse header row
    header_cells = parse_table_row(table_lines[0])
    num_cols = len(header_cells)
    
    # Find data rows (skip separator line)
    data_rows = []
    for line in table_lines[1:]:
        if not is_table_separator(line):
            cells = parse_table_row(line)
            if len(cells) == num_cols:
                data_rows.append(cells)
            elif len(cells) > 0:
                # Pad or truncate to match column count
                cells = cells[:num_cols] + [''] * max(0, num_cols - len(cells))
                data_rows.append(cells)
    
    # Create Word table
    num_rows = 1 + len(data_rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill header row
    for col_idx, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[col_idx]
        # Clean bold markers
        cell_text = cell_text.replace('**', '')
        para = cell.paragraphs[0]
        run = para.add_run(cell_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set header background
        set_cell_shading(cell, '003057')  # Navy
    
    # Fill data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            # Clean bold markers
            cell_text = cell_text.replace('**', '')
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Alternating row colors
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'f4f6f8')  # Light grey
    
    # Add spacing after table
    doc.add_paragraph()

def parse_markdown_line(line):
    """Determine the type and content of a markdown line."""
    line = line.strip()
    
    if not line:
        return ('empty', '')
    
    # Table row detection
    if is_table_row(line):
        return ('table_row', line)
    
    # Headers with # prefix
    if line.startswith('# '):
        return ('h1', line[2:])
    if line.startswith('## '):
        return ('h2', line[3:])
    if line.startswith('### '):
        return ('h3', line[4:])
    if line.startswith('#### '):
        return ('h4', line[5:])
    
    # Roman numeral section headings (without # prefix)
    # Pattern: I. Title or II. Title or III. Title etc (Part headings)
    part_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.\s+(.+)$', line)
    if part_match and len(line) < 80:
        return ('h1', line)
    
    # Pattern: I.1 Title or II.3 Title (Major section)
    section_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.?\s+(.+)$', line)
    if section_match and len(line) < 100:
        return ('h2', line)
    
    # Pattern: I.1.2 Title or III.2.3 Title (Subsection)
    subsection_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.(\d+)\.?\s+(.+)$', line)
    if subsection_match and len(line) < 120:
        return ('h3', line)
    
    # Pattern: I.1.2.3 Title (Sub-subsection)
    subsubsection_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$', line)
    if subsubsection_match and len(line) < 150:
        return ('h4', line)
    
    # Images (Markdown format)
    img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
    if img_match:
        return ('image', (img_match.group(1), img_match.group(2)))
    
    # Horizontal rule
    if line.startswith('---'):
        return ('hr', '')
    
    # Bullet points
    if line.startswith('- ') or line.startswith('* '):
        return ('bullet', line[2:])
    
    # Numbered list
    num_match = re.match(r'^\d+\.\s+(.+)', line)
    if num_match:
        return ('numbered', num_match.group(1))
    
    # Blockquote
    if line.startswith('> '):
        return ('quote', line[2:])
    
    # Skip INSERT FIGURE placeholders
    if '[INSERT FIGURE' in line:
        return ('skip', '')
    
    # Default: body text
    return ('body', line)

def add_image(doc, caption, path):
    """Add an image to the document."""
    # Resolve path
    if path.startswith('figures/'):
        full_path = os.path.join(FIGURES_DIR, os.path.basename(path))
    else:
        full_path = os.path.join(BASE_DIR, path)
    
    if not os.path.exists(full_path):
        print(f"WARNING: Image not found: {full_path}")
        p = doc.add_paragraph(f"[Figure: {caption}]", style='WP Caption')
        return
    
    # Add image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full_path, width=Inches(5.5))
    
    # Add caption
    if caption:
        p = doc.add_paragraph(caption, style='WP Caption')

def process_markdown_to_docx(doc, md_content):
    """Convert markdown content to styled DOCX elements."""
    lines = md_content.split('\n')
    
    # Track which figures have been inserted
    inserted_figures = set()
    
    # Figure trigger keywords and their corresponding files
    FIGURE_TRIGGERS = {
        'Household Pet Ownership': ('Figure1_Pet_Ownership.png', 'Figure 1: Household Pet Ownership Rates by Region'),
        'Pet Population in Europe': ('Figure2_EU_Pet_Pop.png', 'Figure 2: Pet Population in Europe by Species'),
        'Growth of Cat and Dog': ('Figure3_EU_Growth.png', 'Figure 3: Cat and Dog Population Growth in Europe'),
        'Regional Market Size': ('Figure4_Regional_Market.png', 'Figure 4: Regional Market Size for Pet Nutraceuticals'),
        'Feed Probiotics Market': ('Figure5_Probiotics_Share.png', 'Figure 5: Feed Probiotics Market by Species'),
        'Poultry Production': ('Figure6_Poultry_HPAI.png', 'Figure 6: Poultry Production vs HPAI Impact'),
        'Swine Herd Decline': ('Figure7_Swine_Decline.png', 'Figure 7: European Swine Herd Contraction'),
        'De-Ruminization': ('Figure9_Livestock_Trends.png', 'Figure 8: Livestock Production Trends (De-Ruminization)'),
        'Aquaculture vs Capture': ('Figure10_Aqua_v_Capture.png', 'Figure 9: Aquaculture Overtakes Capture Fisheries'),
        'Format Popularity': ('Figure11_Formats.png', 'Figure 10: Nutraceutical Format Preferences by Species'),
        'Preventive Health Wallet': ('Figure12_Wallet.png', 'Figure 11: Preventive Health Spending Allocation'),
        'Consumer Segmentation': ('Figure13_Segmentation.png', 'Figure 12: Consumer Segmentation by Willingness-to-Pay'),
        'Psychological Factors': ('Figure14_Psychology.png', 'Figure 13: Psychological Drivers of Purchase Decisions'),
        'Mobility Supplement Market Evolution': ('Figure15_Mobility_Evo.png', 'Figure 14: Mobility Supplement Premiumization'),
        'Senior': ('Figure16_Senior_Growth.png', 'Figure 15: Senior Pet Market Growth'),
        'Value Chain': ('Figure17_Value_Chain.png', 'Figure 16: Value Chain Economics'),
        'Channel Economics': ('Figure18_Channel_Economics.png', 'Figure 17: Channel Economics Over Time'),
        'Livestock Premix': ('Figure19_Value_Waterfall.png', 'Figure 18: Value Capture Comparison'),
        'Risk/Reward': ('Figure20_Risk_Reward.png', 'Figure 19: Risk/Reward Segment Map'),
        'Pharma Encroachment': ('Figure21_Pharma_Funnel.png', 'Figure 20: Pharma Encroachment Funnel'),
        'Regulatory Landscape': ('Table_US_vs_EU.png', 'Table 1: US vs EU Regulatory Comparison'),
        'Regulatory Timeline': ('Timeline_Regulations.png', 'Figure 21: Regulatory Timeline'),
        'Species Matrix': ('Matrix_Species_Functional.png', 'Figure 22: Species vs Functional Needs Matrix'),
        'Efficacy Matrix': ('Matrix_Efficacy.png', 'Figure 23: Ingredient Efficacy Landscape'),
    }
    
    current_bullet_list = []
    in_bullet = False
    
    # Table row collection
    current_table_rows = []
    in_table = False
    
    for line in lines:
        line_type, content = parse_markdown_line(line)
        
        # Handle table row collection
        if in_table and line_type != 'table_row':
            # End of table, render it
            if current_table_rows:
                add_markdown_table(doc, current_table_rows)
            current_table_rows = []
            in_table = False
        
        # Handle bullet list grouping
        if in_bullet and line_type != 'bullet':
            for bullet in current_bullet_list:
                p = doc.add_paragraph(f"• {bullet}", style='WP Bullet')
            current_bullet_list = []
            in_bullet = False
        
        # Check if we should insert a figure after this line
        for trigger, (fig_file, fig_caption) in FIGURE_TRIGGERS.items():
            if trigger.lower() in line.lower() and fig_file not in inserted_figures:
                # Insert figure after processing this line
                pass  # Will handle below
        
        if line_type == 'empty':
            continue
        elif line_type == 'skip':
            continue
        elif line_type == 'h1':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = COLORS['accent']
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
        elif line_type == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = COLORS['primary']
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            # Check for figure insertion after h2
            for trigger, (fig_file, fig_caption) in FIGURE_TRIGGERS.items():
                if trigger.lower() in content.lower() and fig_file not in inserted_figures:
                    add_image(doc, fig_caption, f"figures/{fig_file}")
                    inserted_figures.add(fig_file)
                    break
        elif line_type == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = COLORS['secondary']
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line_type == 'h4':
            p = doc.add_paragraph()
            run = p.add_run(content.upper())
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = COLORS['grey']
            p.paragraph_format.space_before = Pt(12)
        elif line_type == 'image':
            caption, path = content
            add_image(doc, caption, path)
        elif line_type == 'hr':
            add_page_break(doc)
        elif line_type == 'bullet':
            in_bullet = True
            current_bullet_list.append(content)
        elif line_type == 'numbered':
            p = doc.add_paragraph(f"• {content}", style='WP Bullet')
        elif line_type == 'quote':
            doc.add_paragraph(content, style='WP Quote')
        elif line_type == 'table_row':
            in_table = True
            current_table_rows.append(content)
        elif line_type == 'body':
            text = content
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\[\d+\]', '', text)
            doc.add_paragraph(text, style='WP Body')
    
    # Flush any remaining tables
    if current_table_rows:
        add_markdown_table(doc, current_table_rows)
    
    # Flush any remaining bullets
    if current_bullet_list:
        for bullet in current_bullet_list:
            p = doc.add_paragraph(f"• {bullet}", style='WP Bullet')
    
    # Add any remaining figures that weren't inserted contextually
    # Insert them at the end as an appendix
    remaining_figures = set(FIGURE_TRIGGERS.values()) - {(f, c) for f, c in FIGURE_TRIGGERS.values() if f in inserted_figures}
    if remaining_figures:
        add_page_break(doc)
        doc.add_paragraph("APPENDIX: Additional Figures", style='WP Heading 1')
        for fig_file, fig_caption in remaining_figures:
            add_image(doc, fig_caption, f"figures/{fig_file}")

def main():
    print("=== Generating Publication-Quality DOCX ===")
    
    # Create document
    doc = Document()
    
    # Create custom styles
    print("Creating styles...")
    create_styles(doc)
    
    # Add cover page
    print("Adding cover page...")
    add_cover_page(doc)
    
    # Add executive summary
    print("Adding executive summary...")
    add_executive_summary(doc)
    
    # Add table of contents
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    # Read and process markdown content
    print("Processing main content...")
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    process_markdown_to_docx(doc, md_content)
    
    # Save document
    print(f"Saving to: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    
    print("=== COMPLETE ===")
    print(f"Output: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
